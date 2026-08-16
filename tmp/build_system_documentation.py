from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path
from datetime import date

ROOT = Path(r"E:\Capstone\TrashTruck")
OUT = ROOT / "output" / "documents"
OUT.mkdir(parents=True, exist_ok=True)
OUTFILE = OUT / "TrashTrack_System_Documentation.docx"

NAVY = "173B2C"
GREEN = "2E6B4F"
MINT = "E7F1EB"
LIGHT = "F3F6F4"
GRAY = "5F6B65"
INK = "17211C"
WHITE = "FFFFFF"
GOLD = "B7832F"
RED = "A23B3B"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(.82); sec.bottom_margin = Inches(.78)
sec.left_margin = Inches(.9); sec.right_margin = Inches(.9)
sec.header_distance = Inches(.35); sec.footer_distance = Inches(.35)

def font(run, size=10.5, bold=False, color=INK, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [
    ("Title", 28, NAVY, 0, 10), ("Subtitle", 13, GRAY, 0, 8),
    ("Heading 1", 18, NAVY, 16, 8), ("Heading 2", 14, GREEN, 12, 6),
    ("Heading 3", 11.5, NAVY, 9, 4)]:
    st = styles[name]; st.font.name = "Aptos Display" if "Heading" in name or name=="Title" else "Aptos"
    st.font.size = Pt(size); st.font.bold = name != "Subtitle"; st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for sname in ["Doc Kicker", "Status Complete", "Status Partial", "Status Planned", "Table Note"]:
    if sname not in styles:
        styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
styles["Doc Kicker"].font.name="Aptos"; styles["Doc Kicker"].font.size=Pt(9); styles["Doc Kicker"].font.bold=True; styles["Doc Kicker"].font.color.rgb=RGBColor.from_string(GREEN)
styles["Doc Kicker"].paragraph_format.space_after=Pt(10)
styles["Table Note"].font.name="Aptos"; styles["Table Note"].font.size=Pt(8.5); styles["Table Note"].font.italic=True; styles["Table Note"].font.color.rgb=RGBColor.from_string(GRAY)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=90, start=110, bottom=90, end=110):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for tag,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        x=tcMar.find(qn("w:"+tag))
        if x is None: x=OxmlElement("w:"+tag); tcMar.append(x)
        x.set(qn("w:w"),str(val)); x.set(qn("w:type"),"dxa")

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement("w:tblHeader"); el.set(qn("w:val"),"true"); trPr.append(el)

def table(headers, rows, widths, small=False):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,(h,w) in enumerate(zip(headers,widths)):
        c=t.rows[0].cells[i]; c.width=Inches(w); shade(c,GREEN); margins(c)
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
        font(p.add_run(h),9 if small else 9.5,True,WHITE)
    set_repeat_table_header(t.rows[0])
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,(val,w) in enumerate(zip(row,widths)):
            c=cells[i]; c.width=Inches(w); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx%2: shade(c,"F7F9F8")
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
            font(p.add_run(str(val)),8.3 if small else 9.1,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def add_bullet(text, level=0):
    p=doc.add_paragraph(style="List Bullet" if level==0 else "List Bullet 2")
    p.paragraph_format.left_indent=Inches(.28+.22*level); p.paragraph_format.first_line_indent=Inches(-.16)
    p.paragraph_format.space_after=Pt(3); font(p.add_run(text),10.2)
    return p

def add_number(text):
    p=doc.add_paragraph(style="List Number"); p.paragraph_format.left_indent=Inches(.32); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(4)
    font(p.add_run(text),10.2); return p

def para(text, boldlead=None):
    p=doc.add_paragraph(); p.paragraph_format.widow_control=True
    if boldlead and text.startswith(boldlead):
        font(p.add_run(boldlead),10.5,True); font(p.add_run(text[len(boldlead):]),10.5)
    else: font(p.add_run(text),10.5)
    return p

def callout(title, body, fill=MINT, accent=GREEN):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    c=t.cell(0,0); c.width=Inches(6.58); shade(c,fill); margins(c,150,180,150,180)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(4); font(p.add_run(title.upper()),9,True,accent)
    p=c.add_paragraph(); p.paragraph_format.space_after=Pt(0); font(p.add_run(body),10.2,False,INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def page_break(): doc.add_page_break()

def add_page_field(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    font(p.add_run("TRASHTRACK  |  SYSTEM DOCUMENTATION   "),8,True,GRAY)
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

# Header/footer
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("TRASHTRACK  /  DANAO CITY"),8,True,GRAY)
add_page_field(sec.footer.paragraphs[0])

# Cover
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.style="Doc Kicker"; p.add_run("CAPSTONE SYSTEM DOCUMENTATION")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
font(p.add_run("TrashTrack"),34,True,NAVY,name="Aptos Display")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(18)
font(p.add_run("Integrated Digital Waste Management System"),16,False,GREEN)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Enhancing Operational Workflows and Budget Accuracy in City Solid Waste Management Through Digital Intervention"),11.5,False,GRAY,True)
doc.add_paragraph()
callout("Documentation basis", "Prepared from the August 10, 2026 source-code review and the capstone manuscript. Implementation statements describe what is observable in the repository; incomplete research objectives are explicitly identified.")
for _ in range(4): doc.add_paragraph()
for line,b in [("Prepared by",True),("Luthar James Jimenez | Antonette Lariosa | Vince Louie Racoma",False),("Bachelor of Science in Information Technology",False),("Southwestern University - PHINMA",False),("Version 1.0 | August 10, 2026",False)]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3); font(p.add_run(line),10.5,b,NAVY if b else GRAY)

page_break()
doc.add_heading("Document Control",1)
table(["Field","Value"],[
    ("Document title","TrashTrack System Documentation"),("Version","1.0"),("Review date","August 10, 2026"),
    ("System type","Cross-platform resident and driver mobile application with web-based administrative portals"),
    ("Primary deployment stack","Expo / React Native, Firebase Authentication and Firestore, Cloudinary, Firebase Cloud Functions"),
    ("Document status","Implementation-based baseline for review, testing, and capstone defense")], [1.65,4.93])
doc.add_heading("Purpose and Intended Audience",2)
para("This document records the current TrashTrack system architecture, implemented capabilities, operational workflows, data structures, controls, deployment requirements, and known limitations. It is intended for the capstone panel, project advisers, developers, CENRO and barangay administrators, DICT reviewers, drivers, and future maintainers.")
doc.add_heading("Status Legend",2)
table(["Status","Meaning"],[
    ("Implemented","A working code path and user interface are present in the reviewed repository."),
    ("Partial / prototype","A user interface or basic workflow exists, but automation, validation, integration, or production hardening is incomplete."),
    ("Planned / not evidenced","The manuscript specifies the feature, but a working implementation was not found in the reviewed repository.")],[1.55,5.03])
callout("Important qualification", "Repository inspection confirms implementation presence, not successful production deployment or field acceptance. Live Firebase, Cloudinary, notification, maps, and AI behavior still depend on valid credentials, services, permissions, datasets, and end-to-end testing.", "FFF4DD", GOLD)

page_break()
doc.add_heading("Table of Contents",1)
toc=[("1. Executive Summary","4"),("2. System Overview","5"),("3. Stakeholders and Access Roles","6"),("4. System Architecture","7"),("5. Functional Modules","9"),("6. Core Operational Workflows","13"),("7. Data Model and Synchronization","16"),("8. Security and Privacy","18"),("9. Installation and Configuration","20"),("10. Operating Guide","22"),("11. Testing and Quality Status","25"),("12. Requirements Traceability and Gap Analysis","27"),("13. Maintenance and Roadmap","30"),("Appendices","32")]
for title,n in toc:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.tab_stops.add_tab_stop(Inches(6.2)); font(p.add_run(title),10.5,False,INK); font(p.add_run("\t"+n),10.5,True,GREEN)

page_break(); doc.add_heading("1. Executive Summary",1)
para("TrashTrack is a multi-role digital waste-management platform designed for Danao City. The reviewed repository implements a shared Expo and React Native application that supports resident reporting and schedule viewing, a driver operations suite, a CENRO administrative dashboard, and a limited DICT governance portal. Firebase provides authentication and real-time data synchronization, while Cloudinary is used for image evidence.")
callout("Current baseline", "The strongest implemented area is operational coordination: residents can submit evidence-based reports; administrators can review reports, create schedules, assign drivers and trucks, and publish announcements; drivers can receive work, complete pickups, or report issues with photos.")
doc.add_heading("Implementation Summary",2)
table(["Capability area","Current assessment","Evidence in repository"],[
    ("Resident services","Implemented","Authentication, home, schedules, reports, announcements, profile, feedback, notifications"),
    ("Driver operations","Implemented with exceptions","Truck selection, dispatch view, completion/issue evidence, schedule and history; a separate driver login screen bypasses authentication"),
    ("CENRO administration","Implemented / partial","Reports, schedules, fleet, drivers, announcements, feedback, overrides, route dispatch, analytics UI"),
    ("DICT governance","Partial","Identity/access and rewards are present; several sidebar destinations explicitly remain under construction"),
    ("AI assistant and image analysis","Partial","Gemini-facing services and analysis fields exist; fallback/mock behavior is present when configuration is missing"),
    ("LSTM forecasting and budget prediction","Planned / not evidenced","No Python, TensorFlow, trained model, forecast service, dataset pipeline, or validation artifacts found"),
    ("Automatic routing, heat map, and geofencing","Partial / planned","Manual report selection and ordered dispatch exist; production optimization, heat-map rendering, ETA/geofence automation are not evidenced")],[1.45,1.45,3.68],True)
para("The present system should be presented as an operational prototype with substantial end-to-end CRUD and real-time coordination features, not yet as a fully validated AI-driven decision-support platform.")

page_break(); doc.add_heading("2. System Overview",1)
doc.add_heading("2.1 Problem Addressed",2)
para("The project responds to fragmented manual scheduling, limited visibility into collection status, weak communication among residents, drivers, and administrators, and insufficient operational data for resource planning. TrashTrack centralizes reports, collection schedules, dispatch status, photo evidence, announcements, feedback, and fleet records.")
doc.add_heading("2.2 System Objectives",2)
for x in ["Provide residents with localized schedules, announcements, notifications, and geotagged reporting.","Give drivers an assigned-work view with proof-of-service and issue reporting.","Give CENRO personnel a web dashboard for reports, fleet, staff, schedules, dispatch, feedback, announcements, and operational controls.","Provide DICT or oversight personnel with access-governance and audit-oriented capabilities.","Build a reliable operational dataset that can later support forecasting, route optimization, and budget analysis."]:
    add_bullet(x)
doc.add_heading("2.3 Current Technical Scope",2)
table(["Layer","Technology","Responsibility"],[
    ("Client","React Native 0.81 / React 19 / Expo 54","Mobile and responsive web interfaces with Expo Router navigation"),
    ("Identity","Firebase Authentication","Email/password and configured social sign-in; role checks use profile data and claims"),
    ("Operational data","Cloud Firestore","Real-time reports, schedules, users, trucks, locations, notifications, feedback, and configuration"),
    ("Media","Cloudinary","Report, driver, completion, issue, feedback, and profile images"),
    ("Automation","Firebase Cloud Functions / n8n hooks","Server-side integration and AI/chat automation support"),
    ("Device services","Expo Location and Notifications","Coordinates, location permissions, reminders, and local notifications"),
    ("AI","Google Generative AI integration","Waste-related assistant and image-oriented analysis; requires API configuration")],[1.1,2.0,3.48],True)
doc.add_heading("2.4 Explicit Limitations",2)
for x in ["The reviewed repository contains no implemented LSTM training or inference pipeline.","Route optimization is an administrator-selected dispatch ordering workflow, not a verified vehicle-routing optimization engine.","Some dashboard analytics are descriptive or interface-level rather than validated forecasts.","Production behavior depends on external service credentials and deployed Firebase rules/functions.","The repository currently contains lint errors and numerous warnings, so it does not meet a clean static-analysis gate."]:
    add_bullet(x)

page_break(); doc.add_heading("3. Stakeholders and Access Roles",1)
table(["Role","Primary responsibilities","Main access"],[
    ("Resident","Create account, maintain profile, review schedules, submit trash reports, view announcements and notifications, send feedback","Resident mobile tabs and settings"),
    ("Driver","Select assigned truck, start/end shift, receive dispatches, review schedule, complete pickups, report field issues, view history","Driver mobile suite"),
    ("CENRO administrator","Review reports, schedule collections, manage drivers/trucks/barangays, dispatch routes, publish announcements, inspect feedback and analytics","Administrative web dashboard"),
    ("Barangay coordinator","Coordinate localized schedules and records under CENRO governance","Coordinator directory and assigned administrative functions"),
    ("DICT / oversight administrator","Review identity and access, rewards, and future audit/data-management capabilities","DICT portal; partially implemented"),
    ("System maintainer","Configure environments, deploy Firebase resources, monitor integrations, correct defects, back up and migrate data","Repository, Firebase, Cloudinary, and hosting consoles")],[1.05,3.05,2.48],True)
doc.add_heading("3.1 Role Enforcement",2)
para("Role-aware routing and Firestore rules are both used. Client screens check authentication and user profile roles before opening administrative areas. Firestore rules use authentication, ownership, and an admin custom claim for protected writes. Because some rules are intentionally relaxed for development, production deployment requires a final least-privilege review.")
doc.add_heading("3.2 Responsibility Boundaries",2)
for x in ["Residents originate reports but should not administratively approve or modify other users' submissions.","Drivers may update only assigned operational work and should provide required evidence.","CENRO administrators control schedules, fleet, dispatch, announcements, and operational records.","DICT oversight should remain view-oriented except for identity, access, and explicitly authorized governance functions."]:
    add_bullet(x)

page_break(); doc.add_heading("4. System Architecture",1)
doc.add_heading("4.1 Logical Architecture",2)
table(["Experience layer","Application services","Cloud/data layer"],[
    ("Resident mobile application","Authentication, reporting, schedules, notifications, feedback, AI assistant","Firebase Authentication, Firestore, Cloudinary"),
    ("Driver mobile application","Truck assignment, dispatch, completion, issue reporting, history","Firestore listeners and updates, Cloudinary evidence"),
    ("CENRO web dashboard","Reports, schedules, fleet, accounts, route dispatch, analytics, announcements","Firestore administration and real-time listeners"),
    ("DICT portal","Identity/access, rewards, future audit functions","User documents, administrative data, future analytics logs")],[2.0,2.25,2.33],True)
callout("Data-flow model", "User action -> client-side validation -> media upload when required -> Firestore create/update -> real-time listeners update affected resident, driver, or administrator interfaces. Notification scheduling and cloud automation may run alongside the database update.")
doc.add_heading("4.2 Architectural Characteristics",2)
for title,body in [
    ("Cross-platform client.","Expo Router organizes resident, driver, administrator, and DICT routes in one TypeScript project."),
    ("Event-driven synchronization.","Firestore onSnapshot listeners refresh schedules, reports, locations, trucks, announcements, and history without manual page reloads."),
    ("Evidence-first operations.","Cloudinary images support resident reports, driver completion proof, and issue documentation."),
    ("Service abstraction.","Dedicated services handle notifications, locations, images, schedules, storage, AI, and Cloudinary interactions."),
    ("Configuration-driven integrations.","Firebase, social authentication, Cloudinary, and n8n settings are separated into configuration modules and environment variables.")]:
    para(title+" "+body, title+" ")

page_break(); doc.add_heading("4.3 Repository Structure",2)
table(["Path","Purpose"],[
    ("app/","Expo Router pages and role-specific navigation"),("app/(tabs)/","Resident home, schedule, report, profile, and announcement experiences"),
    ("app/(driver)/","Driver dashboard, truck selection, profile, inbox, schedule, and history"),("app/admin/","Administrative authentication and dashboard container"),
    ("app/dict/","DICT portal routing and role verification"),("components/admin/cenro/","CENRO dashboard modules"),
    ("components/admin/dict/","DICT identity/access and rewards modules"),("components/driver/","Driver completion, issue, and profile components"),
    ("services/","Location, notification, schedule notification, image, Cloudinary, and waste-AI services"),("config/","Firebase, social authentication, Cloudinary, and n8n configuration"),
    ("functions/","Firebase Cloud Functions and AI/webhook logic"),("firestore.rules / firestore.indexes.json","Database authorization rules and indexes")],[1.9,4.68])
doc.add_heading("4.4 External Dependencies",2)
para("TrashTrack depends on Firebase project availability, Cloudinary upload configuration, device location and notification permissions, network access, and any configured AI or webhook provider. The application must degrade safely when a provider is unavailable and must never represent mock data as a production result.")

page_break(); doc.add_heading("5. Functional Modules",1)
doc.add_heading("5.1 Resident Module",2)
table(["Function","Behavior","Status"],[
    ("Authentication and profile","Signup/login, verification flow, profile and password management, barangay selection","Implemented"),
    ("Home dashboard","Next collection, quick actions, recent reports, points, announcements, notifications","Implemented; some calculated metrics are mock-derived"),
    ("Schedule calendar","Filters Firestore barangay schedules to the resident's registered barangay","Implemented"),
    ("Trash reporting","Photo, location, barangay/street, title, description, upload, Firestore submission","Implemented"),
    ("Report history","Lists the resident's reports and details/status","Implemented"),
    ("Announcements","Displays published announcements with filters and details","Implemented"),
    ("Feedback","Submits user ratings/comments for administrative review","Implemented"),
    ("AI assistant","Provides system and waste-management guidance using app context","Partial; configured AI or fallback mock response"),
    ("Live proximity alert","Notify based on actual truck geofence and ETA","Not evidenced as complete")],[1.35,3.65,1.58],True)
doc.add_heading("5.2 Resident Report Data",2)
for x in ["Authenticated user identity and profile context","Title and narrative description","Barangay, street/landmark, and coordinate data when available","Photo evidence and Cloudinary reference","Creation time and processing status","Optional AI analysis such as predicted waste type or estimated weight"]:
    add_bullet(x)

page_break(); doc.add_heading("5.3 Driver Module",2)
table(["Function","Behavior","Status"],[
    ("Truck selection","Reads fleet data and assigns an available truck to the current driver","Implemented with availability controls"),
    ("Shift control","Starts operational context and releases truck/user assignment on shift end","Implemented"),
    ("Dispatch dashboard","Receives assigned schedules and orders them by route sequence","Implemented"),
    ("Complete pickup","Requires image evidence; updates the schedule with completion metadata","Implemented"),
    ("Report issue","Requires photo and description; records issue metadata against the schedule","Implemented"),
    ("Schedule and history","Displays pending and completed/issue work","Implemented"),
    ("Live navigation","Map/location components exist, but production turn-by-turn route navigation is not established","Partial"),
    ("Capacity alert","Full operational capacity broadcasting is not evidenced","Planned / partial"),
    ("Driver authentication","Normal authenticated paths exist, but app/driver-login.tsx explicitly bypasses real authentication","Unsafe prototype path")],[1.35,3.65,1.58],True)
doc.add_heading("5.4 Driver Evidence Controls",2)
para("Completion requires a photo before the completion action becomes available. Issue reporting requires both an image and written description. Firestore rules attempt to restrict an assigned driver's changes to a defined set of completion and issue fields.")

page_break(); doc.add_heading("5.5 CENRO Administrative Module",2)
table(["Area","Implemented capability"],[
    ("Dashboard","Operational totals, time filters, and summarized records"),("Trash reports","Review submitted reports and image evidence; select items for routing"),
    ("Service feedback","Review ratings and written resident feedback"),("Fleet inventory","Maintain truck records and status"),
    ("Driver onboarding and directory","Create or manage driver identities and assignments"),("Route dispatch","Select reports, choose a driver, assign route order, and create schedules"),
    ("Collection schedules","Create recurring barangay schedules and date-specific entries; update or delete them"),("Announcements","Create and publish public operational notices"),
    ("Barangays and coordinators","Maintain service-area and coordinator records"),("Operational overrides","Store administrative configuration and exception settings"),
    ("Waste analytics","Loads report information and AI-oriented insights; full LSTM forecasts are not implemented")],[1.95,4.63])
doc.add_heading("5.6 DICT / Oversight Module",2)
para("The DICT dashboard verifies the user's DICT role and exposes navigation for dashboard, rewards, identity and access, data management, fleet operations, and CENRO command. Only rewards and identity/access have dedicated module implementations in the reviewed routing; the remaining destinations display an under-construction placeholder.")

page_break(); doc.add_heading("6. Core Operational Workflows",1)
doc.add_heading("6.1 Resident Trash Report",2)
for x in ["Resident signs in and opens Report.","Application retrieves or requests location and resident profile context.","Resident captures/selects an image and supplies the required report details.","Client validates required fields and uploads the image to Cloudinary when configured.","Application writes the report to the Firestore reports collection.","Resident history and administrative report listeners reflect the new submission.","Administrator reviews the report and may select it for dispatch."]:
    add_number(x)
doc.add_heading("6.2 Schedule Publication",2)
for x in ["Administrator specifies barangay, date/time or recurrence, waste category, and operational assignment information.","Client validates required fields and writes the schedule to barangay_schedules.","Resident schedule listeners filter the record by the resident's barangay.","Home and schedule screens recalculate upcoming collection information.","Notification service may schedule reminders on supported mobile platforms."]:
    add_number(x)
doc.add_heading("6.3 Route Dispatch",2)
for x in ["Administrator reviews pending resident reports.","Administrator selects reports and an available driver.","The route module creates schedule items with an explicit route order.","Selected reports are updated to indicate dispatch or assignment.","The assigned driver's real-time listener receives and orders the new work."]:
    add_number(x)
callout("Accuracy note", "This workflow performs ordered dispatch. It should not be documented as mathematically optimized routing until a route objective, constraints, algorithm, map-distance service, test dataset, and benchmark results are implemented.", "FFF4DD", GOLD)

page_break(); doc.add_heading("6.4 Pickup Completion and Issue Handling",2)
table(["Stage","Successful pickup","Operational issue"],[
    ("Driver input","Completion photo and optional notes","Issue photo, issue type, and required description"),
    ("Validation","Completion action disabled until image is present","Submission disabled until image and description are present"),
    ("Database update","Status becomes completed with time, driver identity, image, and notes","Status becomes issue with time, reporter identity, image, type, and description"),
    ("Administrative effect","Record becomes proof of service and enters history","Record becomes an exception requiring review or rescheduling")],[1.15,2.72,2.72],True)
doc.add_heading("6.5 Announcement and Notification Flow",2)
for x in ["Authorized administrator creates and publishes an announcement.","Authenticated resident clients query published announcements and order/filter them for display.","While supported clients are active, detection of a newly published item can schedule a local announcement notification.","Pickup reminder logic can schedule notifications one day and one hour before a future schedule.","Web notification behavior is intentionally limited compared with native mobile platforms."]:
    add_number(x)
doc.add_heading("6.6 Feedback Flow",2)
para("Residents submit feedback associated with their authenticated identity. Administrative modules retrieve the records, calculate rating distributions, and provide paginated review. Firestore access is intended to limit resident access to owned records while allowing administrators to manage feedback.")

page_break(); doc.add_heading("7. Data Model and Synchronization",1)
doc.add_heading("7.1 Primary Firestore Collections",2)
table(["Collection","Purpose","Typical writers/readers"],[
    ("users","Identity profile, role, barangay, driver/truck context","Owner and administrators"),
    ("reports","Resident trash reports, image/location/status, optional AI analysis","Resident creates; authenticated users read; admin updates"),
    ("barangay_schedules","Published localized recurring and specific schedules","Authenticated admin workflows; public read currently allowed"),
    ("schedules","Driver-assigned work, route order, completion and issue metadata","Admins/owners; assigned drivers limited updates"),
    ("trucks","Fleet identity, availability, assignment, and status","Authenticated users under current development rule"),
    ("truck_locations","Live or recent truck coordinates","Authenticated users under current development rule"),
    ("announcements","Published administrative communications","Admin writes; authenticated reads"),
    ("feedback / feedbacks","Resident ratings and feedback","Owner creates/reads; admin manages"),
    ("userNotifications and notifications","Inbox items and administrative notification records","User/admin behavior varies by collection"),
    ("comments","Announcement comments","Authenticated creator; owner/admin updates"),
    ("barangays","Service-area reference records","Public read; admin write"),
    ("analytics, error_logs, app_config, public_data","Administrative metrics, logs, and configuration","Primarily administrative")],[1.5,3.1,1.98],True)
doc.add_heading("7.2 Synchronization Pattern",2)
para("Firestore snapshot listeners are used extensively, including for announcements, resident schedules, assigned driver work, truck status, resident reports, user profiles, and history. These listeners provide immediate interface updates but also require disciplined unsubscribe handling, indexes, offline/error behavior, and cost monitoring.")

page_break(); doc.add_heading("7.3 Key Data Relationships",2)
table(["Source record","Relationship","Dependent behavior"],[
    ("users.barangay","Matches barangay_schedules.barangayName","Resident receives localized schedule content"),
    ("users current truck reference","Points to trucks document","Driver dashboard reflects assigned fleet asset"),
    ("reports userId","References authenticated resident UID","Ownership, history, and accountability"),
    ("schedules driver fields / driverUid","Associates work with driver","Driver listener and rule-level assignment checks"),
    ("schedules reportIds","Associates dispatch with selected reports","Administrative traceability from report to route"),
    ("media URL and public ID","References Cloudinary resource","Evidence display and potential cleanup"),
    ("notification userId","Targets an authenticated user","Inbox ownership and read status")],[1.65,2.25,2.68],True)
doc.add_heading("7.4 Data Quality Requirements",2)
for x in ["Store server-generated timestamps for authoritative ordering and audit trails.","Use stable UIDs rather than display names or email addresses for driver assignment.","Standardize collection names; the code and rules currently include singular/plural feedback and multiple notification patterns.","Define enumerations for status, waste category, truck status, issue type, and role.","Validate coordinates, barangay names, and schedule time zones.","Retain AI model version, confidence, input provenance, and human verification when predictive features are added."]:
    add_bullet(x)

page_break(); doc.add_heading("8. Security and Privacy",1)
doc.add_heading("8.1 Existing Controls",2)
for x in ["Firebase Authentication protects most operational reads and writes.","Ownership checks compare authenticated UIDs with record user IDs.","Administrative authorization uses a custom admin claim and role-aware client navigation.","Assigned-driver rules limit schedule completion/issue updates to an approved field list.","Passwords are handled through Firebase Authentication rather than stored directly in Firestore.","Configuration is designed to read service credentials from environment variables."]:
    add_bullet(x)
doc.add_heading("8.2 Security Findings Requiring Resolution",2)
table(["Priority","Finding","Required action"],[
    ("Critical","A standalone driver login path bypasses real authentication","Remove or restrict the mock path before pilot/production use"),
    ("High","barangay_schedules, trucks, and truck_locations currently allow broad authenticated writes or public reads for development","Replace with role- and field-specific production rules"),
    ("High","Driver assignment accepts names or email strings in addition to UID","Normalize assignments to immutable authenticated UID"),
    ("High","Cloudinary unsigned upload presets can be abused if insufficiently restricted","Limit formats, size, folders, transformations, and deletion authority"),
    ("Medium","Multiple overlapping collections increase authorization inconsistency","Consolidate schemas and rules"),
    ("Medium","Logs and UI messages may reveal implementation details","Adopt structured, privacy-safe production logging"),
    ("Medium","No documented retention/deletion schedule","Define retention for photos, location records, reports, and accounts")],[.72,2.58,3.28],True)
doc.add_heading("8.3 Privacy Principles",2)
para("Collect only information required for service delivery. Inform users when location and images are captured, obtain permission, limit access by role, protect minors and bystanders visible in photos, define retention periods, and provide correction or deletion processes consistent with applicable Philippine privacy requirements and institutional policies.")

page_break(); doc.add_heading("9. Installation and Configuration",1)
doc.add_heading("9.1 Prerequisites",2)
for x in ["Supported Node.js runtime and npm dependencies","Expo-compatible Android/iOS device, emulator, or supported web browser","Firebase project with Authentication and Firestore enabled","Cloudinary account and restricted upload preset","Configured OAuth clients if social sign-in is enabled","Notification and location permissions on target mobile devices"]:
    add_bullet(x)
doc.add_heading("9.2 Required Environment Configuration",2)
table(["Configuration group","Representative variables / setup"],[
    ("Firebase","API key, auth domain, project ID, storage bucket, sender ID, application ID"),
    ("Cloudinary","Cloud name, upload preset, and folder strategy"),
    ("Social authentication","Google client IDs and configured provider settings"),
    ("AI / automation","Gemini or other AI key and n8n/webhook URL where used"),
    ("Native capabilities","Android/iOS permission configuration for location, camera, media, and notifications")],[1.75,4.83])
doc.add_heading("9.3 Local Setup",2)
for x in ["Obtain the repository and create a local environment file from the provided example or setup process.","Install JavaScript dependencies using the lockfile.","Insert authorized Firebase, Cloudinary, social-authentication, and AI values.","Deploy Firestore rules and indexes to the intended non-production project first.","Start the Expo development service and open the required target platform.","Test authentication, report upload, schedule synchronization, and driver/admin workflows using dedicated test accounts."]:
    add_number(x)
callout("Configuration hygiene", "Never commit secrets, private service-account files, production API keys, or unrestricted upload credentials. Use separate development, testing, and production projects.", "FDECEC", RED)

page_break(); doc.add_heading("9.4 Deployment Checklist",2)
for x in ["Create and verify separate Firebase environments.","Replace development Firestore rules with least-privilege production rules.","Remove mock login, sample responses, test buttons, and placeholder data.","Run lint, type checks, unit/integration tests, and platform builds with zero blocking errors.","Verify Cloudinary restrictions and image cleanup behavior.","Configure Android/iOS signing, package identifiers, privacy declarations, and permission prompts.","Validate notification behavior on physical devices.","Verify Firestore indexes and performance under realistic record volumes.","Establish monitoring, backups, incident response, and rollback procedures.","Perform user acceptance testing with residents, drivers, and administrators before release."]:
    add_bullet("[ ] "+x)
doc.add_heading("9.5 Production Acceptance Gate",2)
para("Production approval should require evidence that authentication cannot be bypassed, security rules enforce each role, external service failures are handled, essential workflows pass on supported platforms, personal data is protected, and no research claim exceeds the verified behavior of the deployed system.")

page_break(); doc.add_heading("10. Operating Guide",1)
doc.add_heading("10.1 Resident Quick Guide",2)
for x in ["Sign up or log in and complete profile and barangay information.","Open Home to review the next collection and recent activity.","Open Schedule to view barangay-specific collection dates and details.","Open Report, add a clear photo and location details, then submit the incident.","Use My Reports to follow submission history and status.","Read Announcements and device notifications for changes.","Use Settings to update the profile, password, preferences, and submit feedback."]:
    add_number(x)
doc.add_heading("10.2 Driver Quick Guide",2)
for x in ["Authenticate with an approved driver account.","Select an available assigned truck and begin the shift.","Review live dispatches and the next scheduled pickup.","At the assigned location, complete the pickup with photo evidence or report an issue with photo and explanation.","Review history to confirm the update was recorded.","End the shift to release the truck and clear the assignment."]:
    add_number(x)
doc.add_heading("10.3 CENRO Administrator Quick Guide",2)
for x in ["Authenticate through the administrator portal.","Review dashboard indicators and new trash reports.","Maintain trucks, drivers, barangays, and coordinators.","Create collection schedules with correct service area, category, timing, driver, and vehicle.","Select actionable reports, define order, and dispatch them to a driver.","Review completion evidence and operational issues.","Publish announcements and monitor service feedback and configuration overrides."]:
    add_number(x)

page_break(); doc.add_heading("10.4 Troubleshooting Guide",2)
table(["Symptom","Likely cause","Recommended response"],[
    ("User cannot log in","Incorrect credentials, unverified email, disabled provider, missing role profile","Verify Firebase Authentication, profile document, provider settings, and role"),
    ("Schedules do not appear","Barangay mismatch, no published record, listener/index/rules issue","Compare users.barangay with barangayName and inspect Firestore access"),
    ("Photo upload fails","Cloudinary configuration, network, file limits, permission issue","Check preset restrictions, device permission, size/type, and service response"),
    ("Location is missing","Permission denied, emulator unset, GPS disabled","Request permission, enable device location, or configure emulator coordinates"),
    ("Driver sees no assignments","Driver identity mismatch or no assigned schedule","Use UID-based assignment and inspect schedule fields"),
    ("Notifications absent","Web platform limitation, denied permission, expired schedule","Test native device permissions and scheduled-notification state"),
    ("AI gives generic response","Missing API key or fallback behavior","Configure provider and label fallback output clearly"),
    ("Admin write denied","Missing admin custom claim or restrictive rule","Verify token refresh, custom claims, and deployed rules")],[1.35,2.25,2.98],True)

page_break(); doc.add_heading("11. Testing and Quality Status",1)
doc.add_heading("11.1 Recommended Test Levels",2)
table(["Level","Coverage"],[
    ("Unit","Validation helpers, date recurrence, schedule filtering, notification timing, data transformations"),
    ("Component","Forms, disabled states, modal evidence requirements, role-based navigation"),
    ("Integration","Firebase Authentication, Firestore security rules, Cloudinary upload, notification permissions, AI/webhook errors"),
    ("End-to-end","Resident report -> admin review -> dispatch -> driver completion -> resident/admin reflection"),
    ("Security","Unauthorized reads/writes, role escalation, assignment spoofing, upload abuse, secret leakage"),
    ("Performance","Large report/schedule collections, listener cost, images, dashboard pagination"),
    ("Usability","Task completion and System Usability Scale evaluation for each stakeholder role")],[1.25,5.33])
doc.add_heading("11.2 Static Analysis Snapshot",2)
callout("August 10, 2026 lint result", "The repository lint run reported 214 findings: 28 errors and 186 warnings. Blocking errors include unescaped JSX entities and Expo environment-variable destructuring rules; warnings include unused imports/variables and missing React Hook dependencies. This result is a development-quality finding, not proof that runtime workflows fail.", "FFF4DD", GOLD)
doc.add_heading("11.3 Minimum End-to-End Acceptance Scenarios",2)
for x in ["Resident account creation, login, profile completion, logout, and password change","Resident report with valid photo and GPS, plus graceful behavior when each is unavailable","Barangay schedule creation and immediate resident/driver synchronization","Admin report review, selection, dispatch, and status update","Driver completion and issue reporting with evidence enforcement","Announcement publication and resident display/notification","Role and Firestore rule attempts by unauthorized users","Offline, slow-network, and external-provider failure behavior"]:
    add_bullet(x)

page_break(); doc.add_heading("12. Requirements Traceability and Gap Analysis",1)
table(["Manuscript requirement","Repository status","Assessment / next evidence"],[
    ("Resident photo + GPS reporting","Implemented","Run E2E tests and preserve sample audit records"),
    ("Localized real-time schedules","Implemented","Validate all 42 barangays and conflict rules"),
    ("Driver assigned routes and proof of service","Implemented","Replace string identity matching with UID; validate proximity requirements"),
    ("CENRO decision-support dashboard","Implemented / partial","Separate descriptive metrics from predictive metrics"),
    ("Real-time synchronization across roles","Implemented","Load-test listeners and verify security rules"),
    ("Announcements, feedback, and notifications","Implemented / partial","Verify push delivery on physical devices and consolidate schemas"),
    ("EnviroHero rewards","Prototype","Replace report-count and fixed-weight mock calculations with verified rules"),
    ("AI image validation / CNN","Partial","Document actual model/provider, confidence, validation dataset, and failure handling"),
    ("Demand-responsive route optimization","Partial","Implement routing objective, distance matrix, constraints, algorithm, and benchmark"),
    ("Live GPS, ETA, geofence proximity alerts","Partial / not evidenced","Implement location publishing, geofences, privacy controls, and device tests"),
    ("LSTM waste forecast","Not evidenced","Build data pipeline, training/inference service, versioned model, evaluation"),
    ("AI-based budget projection","Not evidenced","Map validated forecast outputs to transparent fiscal assumptions"),
    ("Heat map of waste density","Not evidenced as production feature","Implement geospatial aggregation and dashboard visualization"),
    ("Mayoral/DICT audit portal","Partial","Complete audit dashboard, logs, analytics, and view-only authorization")],[2.0,1.2,3.38],True)

page_break(); doc.add_heading("12.1 Highest-Priority Gaps",2)
for title,body in [
    ("1. Authentication integrity.","Remove the driver authentication bypass and verify all privileged routes and database actions."),
    ("2. Production security rules.","Replace broad development access with role-specific rules and automated emulator tests."),
    ("3. Claims-to-code alignment.","Revise manuscript wording so LSTM, CNN, optimization, geofencing, heat maps, and budget forecasting are described as partial or future work until demonstrated."),
    ("4. AI implementation evidence.","Add datasets, preprocessing, model code, saved artifacts, metrics, baselines, model cards, and reproducible evaluation."),
    ("5. Quality gate.","Resolve blocking lint errors, reduce warnings, add type/test scripts, and maintain a repeatable acceptance suite."),
    ("6. Schema normalization.","Unify duplicated feedback and notification collections and standardize identifiers/status values."),
    ("7. Operational resilience.","Document backups, image lifecycle, error monitoring, provider outages, and rollback.")]:
    para(title+" "+body,title+" ")
doc.add_heading("12.2 Recommended Defense Position",2)
callout("Suggested wording", "TrashTrack currently demonstrates a functional multi-role operational coordination prototype. Its resident reporting, scheduling, administrative dispatch, driver evidence, announcements, and real-time database synchronization are implemented. Advanced predictive forecasting, budget decision support, automatic vehicle routing, geofencing, and audit analytics remain development and validation work.")

page_break(); doc.add_heading("13. Maintenance and Roadmap",1)
table(["Phase","Priority outcomes","Completion evidence"],[
    ("Stabilize","Remove bypasses/mocks, fix lint errors, normalize schemas, harden rules","Clean lint/type run, security emulator tests, reviewed schema"),
    ("Validate operations","Run complete resident-admin-driver workflows on physical devices","Signed UAT scripts, defect log, screenshots, audit records"),
    ("Complete geospatial layer","Reliable truck location, ETA, proximity alerts, map routes","Permission/privacy tests, location accuracy and battery results"),
    ("Implement optimization","Vehicle-routing constraints and measurable routing baseline","Algorithm code, benchmark dataset, distance/fuel comparison"),
    ("Implement forecasting","Historical data pipeline and LSTM/baseline models","Versioned model, RMSE/MAE results, holdout evaluation"),
    ("Add budget decision support","Transparent conversion of forecasted tonnage to resource/budget scenarios","Documented assumptions, administrator validation"),
    ("Govern and deploy","Monitoring, retention, backups, incident response, production rollout","Runbooks, approvals, deployment and rollback records")],[1.05,3.05,2.48],True)
doc.add_heading("13.1 Maintenance Cadence",2)
for x in ["Weekly: review error logs, failed uploads, notification failures, and security anomalies.","Monthly: review dependency updates, Firestore costs/indexes, stale accounts, storage growth, and backup restoration.","Per release: run lint, type checks, automated tests, security-rule tests, platform builds, and regression scenarios.","Quarterly: review privacy retention, role assignments, external-provider keys, disaster recovery, and user feedback.","Per AI model release: retain dataset version, preprocessing, hyperparameters, metrics, limitations, approval, and rollback model."]:
    add_bullet(x)

page_break(); doc.add_heading("Appendix A. Functional Inventory",1)
table(["Interface / component","Role","Purpose"],[
    ("Resident Home","Resident","Next schedule, recent activity, quick actions, announcements/notifications context"),("Schedule","Resident","Barangay-filtered calendar and collection details"),
    ("Report","Resident","Evidence-based trash report with location and description"),("My Reports","Resident","Personal report history and details"),
    ("Announcements","Resident","Published administrative notices"),("Settings/Profile","Resident","Profile, password, theme/settings, and feedback"),
    ("Driver Dashboard","Driver","Shift state, dispatches, next pickup, completion/issue actions"),("Select Truck","Driver","Fleet selection and assignment"),
    ("Driver Schedule/History","Driver","Pending and resolved operational tasks"),("CENRO Dashboard","Administrator","Operational overview and module navigation"),
    ("Trash Reports","Administrator","Review resident reports and evidence"),("Route Optimization","Administrator","Select, order, and dispatch reports"),
    ("Collection Scheduler","Administrator","Recurring and specific barangay schedules"),("Fleet / Drivers","Administrator","Truck and personnel administration"),
    ("Announcements / Feedback","Administrator","Public communication and service review"),("DICT Identity & Access","DICT","Review/manage user access according to implemented controls"),
    ("DICT Rewards","DICT","Review reward/token-oriented records"),("AI Chat","Resident","Context-aware system and waste guidance")],[2.4,1.0,3.18],True)

page_break(); doc.add_heading("Appendix B. Key Repository Evidence",1)
table(["Evidence type","Representative locations"],[
    ("Resident pages","app/(tabs)/home.tsx, schedule.tsx, report.tsx, profile.tsx, announcements.tsx; app/my-reports.tsx"),
    ("Driver pages","app/(driver)/index.tsx, select-truck.tsx, inbox.tsx, pages/DriverSchedulePage.tsx, pages/DriverHistoryPage.tsx"),
    ("Administrative pages","app/admin/dashboard.tsx and components/admin/cenro/*.tsx"),
    ("DICT pages","app/dict/dashboard.tsx and components/admin/dict/*.tsx"),
    ("Operational services","services/locationService.ts, notificationService.ts, scheduleNotificationService.ts, cloudinaryService.ts, wasteAIService.ts"),
    ("Configuration","config/firebase.ts, cloudinary.ts, socialAuth.ts, n8n.ts"),
    ("Backend and policy","functions/index.js, firestore.rules, firestore.indexes.json, firebase.json"),
    ("Project manifest","package.json and package-lock.json")],[1.72,4.86])
doc.add_heading("Appendix C. Source Basis",1)
para("This documentation was developed from the TrashTrack repository at E:\\Capstone\\TrashTruck and the capstone manuscript titled “Enhancing Operational Workflows and Budget Accuracy in City Solid Waste Management Through Digital Intervention” (Danao city newer.pdf). The manuscript provided objectives and intended features; repository inspection determined the implementation status recorded here.")
doc.add_heading("Appendix D. Approval Record",1)
table(["Role","Name","Signature","Date"],[
    ("Project lead","","",""),("Technical adviser","","",""),("CENRO representative","","",""),("Capstone panel / reviewer","","","")],[1.55,1.85,1.85,1.33])

# Document metadata
doc.core_properties.title = "TrashTrack System Documentation"
doc.core_properties.subject = "Implementation-based technical and operational documentation"
doc.core_properties.author = "TrashTrack Capstone Team"
doc.core_properties.keywords = "TrashTrack, Danao City, waste management, system documentation"
doc.core_properties.comments = "Generated from reviewed repository and capstone manuscript; implementation statuses must be revalidated after major changes."

# avoid row splitting
for t in doc.tables:
    for row in t.rows:
        trPr=row._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); trPr.append(cant)
        for cell in row.cells:
            for p in cell.paragraphs: p.paragraph_format.keep_together=True

doc.save(OUTFILE)
print(OUTFILE)
